"""
Extract raw text from PDF and DOCX files.

Works with both local filesystem and S3/cloud storage backends:
  - Never calls .path on a FieldFile (raises NotImplementedError on S3)
  - Uses .open() + a temp file so any storage backend is supported
"""
import os
import logging
import tempfile
from pathlib import Path

logger = logging.getLogger(__name__)


# ── Low-level extractors (accept a local file path) ──────────────────────────

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a local PDF file using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        text_parts = []
        with fitz.open(file_path) as doc:
            for page in doc:
                text_parts.append(page.get_text("text"))
        return "\n".join(text_parts).strip()
    except Exception as e:
        logger.error(f"PDF extraction failed for {file_path}: {e}")
        raise ValueError(f"Could not extract text from PDF: {e}") from e


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a local DOCX file."""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also grab table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs).strip()
    except Exception as e:
        logger.error(f"DOCX extraction failed for {file_path}: {e}")
        raise ValueError(f"Could not extract text from DOCX: {e}") from e


def extract_text(file_path: str, file_type: str) -> str:
    """Route to the correct parser by extension (local path version)."""
    ext = file_type.lower().lstrip(".")
    if ext == "pdf":
        return extract_text_from_pdf(file_path)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(file_path)
    raise ValueError(f"Unsupported file type: {file_type}")


# ── S3-safe entry point (used by Celery tasks) ────────────────────────────────

def extract_text_from_field(file_field, file_type: str) -> str:
    """
    Extract text from a Django FieldFile regardless of storage backend.

    Works with:
      - Local FileSystemStorage  (.path available)
      - AWS S3 / any other cloud storage  (.path raises NotImplementedError)

    Strategy: stream the file into a NamedTemporaryFile, parse it, then delete.
    """
    ext = file_type.lower().lstrip(".")
    suffix = f".{ext}"

    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp_path = tmp.name
            # .open() works on every Django storage backend
            with file_field.open("rb") as src:
                while True:
                    chunk = src.read(8 * 1024 * 1024)  # 8 MB chunks
                    if not chunk:
                        break
                    tmp.write(chunk)

        return extract_text(tmp_path, file_type)

    except Exception as e:
        logger.error(f"extract_text_from_field failed (type={file_type}): {e}")
        raise
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
